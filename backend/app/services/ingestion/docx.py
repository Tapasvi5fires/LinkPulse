import logging
from typing import List
from datetime import datetime
import os
from app.services.ingestion.base import BaseIngestor, IngestedDocument
import docx

logger = logging.getLogger(__name__)

class DocxIngestor(BaseIngestor):
    async def ingest(self, source: str, **kwargs) -> List[IngestedDocument]:
        """
        Ingest a DOCX file from a file path.
        """
        from starlette.concurrency import run_in_threadpool
        return await run_in_threadpool(self._ingest_sync, source)

    def _ingest_sync(self, source: str) -> List[IngestedDocument]:
        documents = []
        try:
            doc = docx.Document(source)
            full_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))

            text_content = "\n\n".join(full_text)
            
            if text_content.strip():
                documents.append(IngestedDocument(
                    content=text_content,
                    source_url=source,
                    source_type="docx",
                    metadata={
                        "title": os.path.basename(source),
                        "ingested_at": datetime.utcnow().isoformat(),
                        "source_url": source, # Inject for persistence
                        "source_type": "docx"
                    }
                ))
        except Exception as e:
            logger.error(f"Error ingesting DOCX {source}: {e}")
            raise e
            
        return documents
